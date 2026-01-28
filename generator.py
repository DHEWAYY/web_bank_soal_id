import json
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- 1. CETAKAN HTML UTAMA ---
TEMPLATE_PG = """<article class="bg-white p-6 rounded-xl shadow-sm border border-gray-100 mb-6"><div class="flex gap-3"><span class="bg-blue-100 text-blue-700 font-bold px-3 py-1 rounded h-fit text-sm">{NO}.</span><div class="w-full"><p class="text-lg font-medium mb-4">{PERTANYAAN}</p><div class="grid grid-cols-1 md:grid-cols-2 gap-3 mb-4"><button class="text-left px-4 py-2 rounded border hover:bg-blue-50 text-sm"><span class="font-bold mr-2">A.</span> {OPSI_A}</button><button class="text-left px-4 py-2 rounded border hover:bg-blue-50 text-sm"><span class="font-bold mr-2">B.</span> {OPSI_B}</button><button class="text-left px-4 py-2 rounded border hover:bg-blue-50 text-sm"><span class="font-bold mr-2">C.</span> {OPSI_C}</button><button class="text-left px-4 py-2 rounded border hover:bg-blue-50 text-sm"><span class="font-bold mr-2">D.</span> {OPSI_D}</button></div><details class="group"><summary class="flex cursor-pointer items-center gap-2 text-blue-600 font-semibold text-sm select-none"><i class="fa-solid fa-key"></i> Lihat Pembahasan</summary><div class="mt-3 bg-gray-50 border-l-4 border-blue-500 p-4 text-sm rounded"><p class="font-bold text-gray-900 mb-1">Jawaban: {JAWABAN}</p><p class="text-gray-700 whitespace-pre-line">{PEMBAHASAN}</p></div></details></div></div></article>"""
TEMPLATE_ESSAY = """<article class="bg-white p-6 rounded-xl shadow-sm border border-orange-100 mb-6"><div class="flex gap-3"><span class="bg-orange-100 text-orange-700 font-bold px-3 py-1 rounded h-fit text-sm">Esai {NO}.</span><div class="w-full"><p class="text-lg font-medium mb-4">{PERTANYAAN}</p><textarea class="w-full border p-3 rounded-lg text-sm mb-3 focus:outline-blue-500" rows="3" placeholder="Tulis jawabanmu disini..."></textarea><details class="group"><summary class="flex cursor-pointer items-center gap-2 text-orange-600 font-semibold text-sm select-none"><i class="fa-solid fa-book-open"></i> Lihat Jawaban Lengkap</summary><div class="mt-3 bg-orange-50 border-l-4 border-orange-500 p-4 text-sm rounded"><p class="font-bold text-gray-900 mb-1">Pembahasan:</p><p class="text-gray-700 whitespace-pre-line font-mono text-xs md:text-sm">{JAWABAN_LENGKAP}</p></div></details></div></div></article>"""

# --- 2. CETAKAN HALAMAN LIST (INDEX) ---
TEMPLATE_INDEX = """
<!DOCTYPE html>
<html lang="id">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{JUDUL_HALAMAN} | BankSoal.id</title>
    <script src="https://cdn.tailwindcss.com"></script>
    <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.4.0/css/all.min.css">
</head>
<body class="bg-gray-50 font-sans text-gray-800">
    <nav class="bg-white border-b shadow-sm sticky top-0 z-50">
        <div class="max-w-4xl mx-auto px-4 h-16 flex items-center justify-between">
            <a href="index.html" class="font-bold text-xl text-blue-600 hover:text-blue-800 transition flex items-center gap-2">
                <i class="fa-solid fa-graduation-cap"></i> BankSoal.id
            </a>
            <div class="flex gap-4 text-sm font-semibold text-gray-500">
                <a href="index_sd.html" class="hover:text-blue-600 transition {AKTIF_SD}">SD</a>
                <a href="index_smp.html" class="hover:text-blue-600 transition {AKTIF_SMP}">SMP</a>
                <a href="index_sma.html" class="hover:text-blue-600 transition {AKTIF_SMA}">SMA</a>
                <a href="index_smk.html" class="hover:text-blue-600 transition {AKTIF_SMK}">SMK</a>
            </div>
        </div>
    </nav>
    
    <main class="max-w-4xl mx-auto px-4 py-8">
        <div class="w-full h-[100px] bg-gray-200 rounded-lg flex items-center justify-center mb-8 border-2 border-dashed border-gray-300">
            <span class="text-gray-500 font-bold text-sm">[IKLAN ADSENSE - DISPLAY]</span>
        </div>

        <div class="text-center mb-10">
            <h1 class="text-3xl font-bold text-gray-900 mb-3">{JUDUL_HEADER}</h1>
            <p class="text-gray-600">Total ada <b>{JUMLAH_SOAL}</b> materi latihan siap dikerjakan.</p>
        </div>

        <div class="grid grid-cols-1 md:grid-cols-2 gap-4 mb-10">
            {LIST_LINK}
        </div>
        
        <div class="w-full h-[250px] bg-gray-200 rounded-lg flex items-center justify-center border-2 border-dashed border-gray-300">
            <span class="text-gray-500 font-bold text-sm">[IKLAN ADSENSE - KOTAK BESAR]</span>
        </div>
    </main>
    <footer class="text-center py-8 text-gray-400 text-sm border-t mt-8">&copy; 2024 BankSoal.id Engine</footer>
</body>
</html>
"""

# --- 3. FUNGSI WORD GENERATOR ---
def create_docx(data, filename_base):
    doc = Document()
    meta = data.get('meta', {})
    doc.add_heading(meta.get('judul_bab', 'Latihan Soal'), 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Mapel: {meta.get('mapel')} | Kelas: {meta.get('kelas')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("-" * 80)
    doc.add_heading('A. Pilihan Ganda', level=1)
    for q in data.get('soal_pg', []):
        doc.add_paragraph(f"{q['no']}. {q['tanya']}")
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.5)
        p.add_run(f"A. {q['opsi_a']}\nB. {q['opsi_b']}\nC. {q['opsi_c']}\nD. {q['opsi_d']}")
    doc.add_heading('B. Soal Uraian', level=1)
    for q in data.get('soal_essay', []): doc.add_paragraph(f"{q['no']}. {q['tanya']}\n")
    doc.add_page_break(); doc.add_heading('KUNCI JAWABAN', level=1)
    for q in data.get('soal_pg', []): doc.add_paragraph().add_run(f"{q['no']}. {q['jawaban']}").bold = True
    for q in data.get('soal_essay', []): doc.add_paragraph().add_run(f"{q['no']}. {q['jawaban_lengkap']}")
    os.makedirs('output/downloads', exist_ok=True)
    doc.save(f"output/downloads/{filename_base}.docx")
    return f"downloads/{filename_base}.docx"

# --- 4. FUNGSI INDEX HALAMAN ---
def create_index_page(filename, title, header, materials, active_menu=""):
    links_html = ""
    for m in materials:
        links_html += f"""<a href="{m['link']}" class="block bg-white border rounded-xl p-5 hover:shadow-md hover:border-blue-300 transition group"><h3 class="font-bold text-lg text-gray-800 group-hover:text-blue-600">{m['judul']}</h3><p class="text-sm text-gray-500 mt-1"><span class="bg-gray-100 text-gray-600 px-2 py-0.5 rounded text-xs font-bold mr-2">{m['jenjang']}</span> {m['info']}</p><div class="mt-3 text-blue-500 text-sm font-semibold flex items-center gap-2">Buka Soal <i class="fa-solid fa-arrow-right"></i></div></a>"""
    
    if not links_html: links_html = '<div class="col-span-2 text-center py-10 text-gray-400">Belum ada materi untuk kategori ini.</div>'

    active_sd = "text-blue-600" if active_menu == "SD" else ""
    active_smp = "text-blue-600" if active_menu == "SMP" else ""
    active_sma = "text-blue-600" if active_menu == "SMA" else ""
    active_smk = "text-blue-600" if active_menu == "SMK" else ""

    html = TEMPLATE_INDEX.replace("{JUDUL_HALAMAN}", title).replace("{JUDUL_HEADER}", header).replace("{JUMLAH_SOAL}", str(len(materials))).replace("{LIST_LINK}", links_html).replace("{AKTIF_SD}", active_sd).replace("{AKTIF_SMP}", active_smp).replace("{AKTIF_SMA}", active_sma).replace("{AKTIF_SMK}", active_smk)
    
    with open(f'output/{filename}', 'w', encoding='utf-8') as f: f.write(html)
    print(f"🏠 Halaman Index Dibuat: {filename}")

# --- 5. FUNGSI UTAMA (ANTI CRASH & STRICT FALSE) ---
def generate_pages():
    try:
        with open('template.html', 'r', encoding='utf-8') as f: template_utama = f.read()
    except: return print("❌ Template gak ada!")

    folder_data = 'data'
    os.makedirs('output', exist_ok=True)
    files = [f for f in os.listdir(folder_data) if f.endswith('.json')]
    
    all_materials = []

    print(f"🚀 Memproses {len(files)} materi...")

    for filename in files:
        path = os.path.join(folder_data, filename)
        
        try:
            with open(path, 'r', encoding='utf-8') as f:
                # 🔥 MAGIC FIX: strict=False MEMBOLEHKAN ENTER DI DALAM TEKS
                data = json.load(f, strict=False)
            
            # PROSES DATA AMAN
            meta = data.get('meta', {})
            judul_seo = meta.get('judul_seo', meta.get('judul_bab', 'Bank Soal')) 
            
            nama_base = filename.replace('.json', '')
            link_docx = create_docx(data, nama_base) 
            
            html_pg = "".join([TEMPLATE_PG.format(NO=q['no'], PERTANYAAN=q['tanya'], OPSI_A=q['opsi_a'], OPSI_B=q['opsi_b'], OPSI_C=q['opsi_c'], OPSI_D=q['opsi_d'], JAWABAN=q['jawaban'], PEMBAHASAN=q['pembahasan']) for q in data.get('soal_pg', [])])
            html_essay = "".join([TEMPLATE_ESSAY.format(NO=q['no'], PERTANYAAN=q['tanya'], JAWABAN_LENGKAP=q['jawaban_lengkap']) for q in data.get('soal_essay', [])])
            
            konten = f'<h2 class="text-xl font-bold text-blue-800 mb-4 border-b pb-2">A. Pilihan Ganda</h2>{html_pg}<div class="w-full flex justify-center my-8"><div class="w-[300px] h-[250px] bg-gray-200 flex items-center justify-center text-gray-500 text-sm font-bold border-2 border-dashed border-gray-300">[IKLAN TENGAH ARTIKEL]</div></div><h2 class="text-xl font-bold text-orange-800 mt-8 mb-4 border-b pb-2">B. Essay</h2>{html_essay}'
            
            halaman = template_utama.replace("{{JUDUL_BAB}}", judul_seo).replace("{{JUDUL_BAB_H1}}", meta.get('judul_bab', '')).replace("{{JENJANG}}", meta.get('jenjang', '')).replace("{{MAPEL}}", meta.get('mapel', '')).replace("{{KELAS}}", meta.get('kelas', '')).replace("{{LINK_DOWNLOAD}}", link_docx).replace("{{LIST_SOAL}}", konten)
            
            with open(f'output/{nama_base}.html', 'w', encoding='utf-8') as f: f.write(halaman)
            
            all_materials.append({
                'judul': meta.get('judul_bab', 'Tanpa Judul'),
                'jenjang': meta.get('jenjang', 'UMUM').upper(), 
                'info': f"{meta.get('mapel')} - {meta.get('kelas')}",
                'link': f"{nama_base}.html"
            })
            print(f"✅ Sukses: {filename}")

        except json.JSONDecodeError as e:
            # KALAU MASIH ERROR PARAH, KITA KASIH TAU BARISNYA
            print(f"❌ ERROR JSON RUSAK: {filename}")
            print(f"   >>> Baris: {e.lineno}, Kolom: {e.colno} (Cek ada karakter aneh/Enter)")
        except Exception as e:
            print(f"⚠️ Error Lain di {filename}: {e}")

    # GENERATE HALAMAN INDEX
    print("\n🏠 Membuat Halaman Index...")
    create_index_page('index.html', 'Bank Soal Lengkap', 'Gudang Bank Soal Terlengkap', all_materials)
    
    # Filter Jenjang
    create_index_page('index_sd.html', 'Bank Soal SD', 'Kumpulan Soal SD', [m for m in all_materials if 'SD' in m['jenjang']], "SD")
    create_index_page('index_smp.html', 'Bank Soal SMP', 'Kumpulan Soal SMP', [m for m in all_materials if 'SMP' in m['jenjang']], "SMP")
    create_index_page('index_sma.html', 'Bank Soal SMA', 'Kumpulan Soal SMA', [m for m in all_materials if 'SMA' in m['jenjang']], "SMA")
    create_index_page('index_smk.html', 'Bank Soal SMK', 'Kumpulan Soal SMK', [m for m in all_materials if 'SMK' in m['jenjang']], "SMK")

    print("🎉 SELESAI! Coba refresh web lu.")

if __name__ == "__main__":
    print("🏁 Mulai menjalankan script...")
    generate_pages()
